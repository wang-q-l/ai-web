#!/usr/bin/env python3
"""
应用完整的表格样式（表头、内容、行高、宽度）
"""
import sys
from docx import Document
from docx.shared import RGBColor, Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 导入样式配置
import os
sys.path.insert(0, os.path.dirname(__file__))
from cover_style_config import TABLE_STYLE

def apply_table_styles(doc_path):
    """应用完整表格样式到所有表格"""
    doc = Document(doc_path)
    header_config = TABLE_STYLE['header']
    content_config = TABLE_STYLE['content']
    row_height_config = TABLE_STYLE['row_height']

    # 解析背景色
    bg_color = header_config['background_color'].lstrip('#')

    table_count = 0
    for table in doc.tables:
        if len(table.rows) == 0:
            continue

        # 判断是否为垂直表格（左列是标签）
        # 垂直表格特征：只有2列，且行数较多
        is_vertical_table = (len(table.columns) == 2 and len(table.rows) > 3)

        # 处理每一行
        for row_idx, row in enumerate(table.rows):
            # 垂直表格不应用表头样式
            if is_vertical_table:
                config = content_config
                row.height = row_height_config['content']
            else:
                is_header = (row_idx == 0)
                config = header_config if is_header else content_config

                # 设置行高
                if is_header:
                    row.height = row_height_config['header']
                else:
                    row.height = row_height_config['content']

            # 处理每个单元格
            for cell in row.cells:
                # 只有非垂直表格的表头才设置背景色
                if not is_vertical_table and row_idx == 0:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), bg_color)
                    cell._element.get_or_add_tcPr().append(shading_elm)

                # 设置字体样式
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = config['font_name']
                        run.font.size = Pt(config['font_size'])
                        run.font.bold = config['bold']

                        # 设置中文字体
                        r = run._element
                        rPr = r.get_or_add_rPr()
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is None:
                            rFonts = OxmlElement('w:rFonts')
                            rPr.insert(0, rFonts)
                        rFonts.set(qn('w:eastAsia'), config['font_name'])

                    # 设置对齐方式
                    if config['alignment'] == 'center':
                        paragraph.alignment = 1
                    elif config['alignment'] == 'right':
                        paragraph.alignment = 2
                    else:
                        paragraph.alignment = 0

        table_count += 1

    doc.save(doc_path)
    print(f'✅ 完整表格样式已应用（共 {table_count} 个表格）')

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python3 apply-table-header-style.py <文档路径>')
        sys.exit(1)

    apply_table_styles(sys.argv[1])
