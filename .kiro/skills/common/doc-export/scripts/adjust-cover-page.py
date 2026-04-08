#!/usr/bin/env python3
"""
为 Word 文档添加封面、目录、页码和表格边框
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys
import os

# 导入样式配置
config_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'config')
sys.path.insert(0, config_dir)
from style_config import TABLE_STYLE

def set_font(run, font_name, font_size):
    """设置字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_table_borders(table):
    """为表格设置边框和行距"""
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # 移除旧的边框设置
    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tblPr.remove(old_borders)

    # 创建新的边框设置
    tblBorders = OxmlElement('w:tblBorders')

    # 设置所有边框：上、下、左、右、内部横线、内部竖线
    for border_name in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  # 边框类型：单线
        border.set(qn('w:sz'), '4')  # 边框粗细：4（1/8磅）
        border.set(qn('w:color'), '000000')  # 边框颜色：黑色
        border.set(qn('w:space'), '0')  # 边框间距
        tblBorders.append(border)

    tblPr.append(tblBorders)

    # 设置单元格内边距（从配置文件读取）
    tblCellMar = tblPr.find(qn('w:tblCellMar'))
    if tblCellMar is None:
        tblCellMar = OxmlElement('w:tblCellMar')
        tblPr.append(tblCellMar)

    # 从配置读取内边距并转换为 twips（1pt = 20 twips）
    cell_margins = TABLE_STYLE['cell_margins']
    margins = {
        'top': str(int(cell_margins['top'] * 20)),
        'bottom': str(int(cell_margins['bottom'] * 20)),
        'left': str(int(cell_margins['left'] * 20)),
        'right': str(int(cell_margins['right'] * 20))
    }

    # 设置上下左右内边距
    for margin_name, value in margins.items():
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), value)  # 单位：twips（1/20磅）
        margin.set(qn('w:type'), 'dxa')
        tblCellMar.append(margin)

    # 设置表格宽度为 100%（根据窗口调整表格）
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), '5000')  # 5000 = 100% (50 = 1%)
    tblW.set(qn('w:type'), 'pct')  # 百分比类型

    # 设置单元格段落行距为 1.0，并取消首行缩进
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.first_line_indent = Pt(0)  # 取消首行缩进

def add_page_number(section):
    """添加页码：第 X 页，共 X 页"""
    footer = section.footer
    footer.is_linked_to_previous = False

    # 清空现有内容
    for para in footer.paragraphs:
        para.clear()

    # 添加页码段落
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加"第 "
    run = para.add_run('第 ')
    set_font(run, '仿宋_GB2312', 10.5)

    # 添加页码域
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._element.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar2)

    run = para.add_run('1')
    set_font(run, '仿宋_GB2312', 10.5)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar3)

    # 添加" 页，共 "
    run = para.add_run(' 页，共 ')
    set_font(run, '仿宋_GB2312', 10.5)

    # 添加总页数域
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar4)

    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = ' NUMPAGES '
    run._element.append(instrText2)

    fldChar5 = OxmlElement('w:fldChar')
    fldChar5.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar5)

    run = para.add_run('1')
    set_font(run, '仿宋_GB2312', 10.5)

    fldChar6 = OxmlElement('w:fldChar')
    fldChar6.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar6)

    # 添加" 页"
    run = para.add_run(' 页')
    set_font(run, '仿宋_GB2312', 10.5)

def generate_toc_content(doc):
    """生成目录内容并插入到TOC域后"""
    # 收集所有标题
    headings = []
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.replace('Heading ', ''))
            if level <= 3:  # 只收集1-3级标题
                headings.append({
                    'text': para.text,
                    'level': level,
                    'para': para
                })

    # 找到TOC域的位置
    toc_index = None
    for i, para in enumerate(doc.paragraphs):
        if 'Table of Contents' in para.text or '目录' in para.text:
            # 检查是否是TOC域
            for run in para.runs:
                if run._element.findall('.//w:fldChar', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    toc_index = i
                    break
            if toc_index is not None:
                break

    if toc_index is None:
        return

    # 在TOC域后插入目录内容
    insert_pos = toc_index + 1
    for heading in headings:
        # 创建目录项段落
        new_para = doc.paragraphs[insert_pos]._element.getparent().insert(
            doc.paragraphs[insert_pos]._element.getparent().index(doc.paragraphs[insert_pos]._element),
            OxmlElement('w:p')
        )

        # 获取新段落对象
        from docx.text.paragraph import Paragraph
        toc_para = Paragraph(new_para, doc)

        # 设置缩进
        indent = (heading['level'] - 1) * 0.5  # 每级缩进0.5cm
        toc_para.paragraph_format.left_indent = Cm(indent)
        toc_para.paragraph_format.first_line_indent = Cm(0)

        # 添加文本
        run = toc_para.add_run(heading['text'])
        set_font(run, '仿宋_GB2312', 14)

        insert_pos += 1

def update_toc_fields(doc):
    """更新目录域，使其自动展开"""
    # 遍历所有段落，查找 TOC 域
    for para in doc.paragraphs:
        for run in para.runs:
            # 查找域代码
            for fldChar in run._element.findall('.//w:fldChar', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                fldCharType = fldChar.get(qn('w:fldCharType'))
                if fldCharType == 'begin':
                    # 设置域为脏状态，强制 Word 打开时更新
                    fldChar.set(qn('w:dirty'), 'true')

def replace_toc_title(doc):
    """将目录标题 'Table of Contents' 替换为 '目录'"""
    for para in doc.paragraphs:
        # 检查段落文本是否包含 Table of Contents
        if 'Table of Contents' in para.text:
            # 清空所有 runs 并重新创建
            for run in para.runs:
                run._element.getparent().remove(run._element)
            # 添加新的文本
            run = para.add_run('目录')
            set_font(run, '仿宋_GB2312', 16)
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break

def remove_manual_toc(doc):
    """删除手动添加的Markdown目录章节"""
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        # 找到"目录"标题
        if para.text.strip() == '目录' and para.style.name.startswith('Heading'):
            # 删除目录标题
            p = para._element
            p.getparent().remove(p)
            # 删除后续的列表项，直到遇到下一个标题
            while i < len(doc.paragraphs):
                next_para = doc.paragraphs[i]
                if next_para.style.name.startswith('Heading'):
                    break
                p = next_para._element
                p.getparent().remove(p)
            break
        i += 1

def insert_cover_page(doc, title, subtitle, author, date):
    """在文档开头插入封面页"""
    # 在第一个段落前插入封面内容
    first_para = doc.paragraphs[0]

    # 标题
    title_para = first_para.insert_paragraph_before(title)
    title_para.style = 'Title'
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 副标题
    subtitle_para = first_para.insert_paragraph_before(subtitle)
    subtitle_para.style = 'Subtitle'
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 空行
    first_para.insert_paragraph_before()
    first_para.insert_paragraph_before()

    # 作者
    author_para = first_para.insert_paragraph_before(author)
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in author_para.runs:
        run.font.name = '仿宋_GB2312'
        run.font.size = Pt(14)

    # 日期
    date_para = first_para.insert_paragraph_before(date)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date_para.runs:
        run.font.name = '仿宋_GB2312'
        run.font.size = Pt(14)

    # 封面后分页
    from docx.enum.text import WD_BREAK
    first_para.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)

def insert_toc(doc):
    """在封面后插入目录"""
    # 找到第一个分页符后的段落
    insert_pos = None
    for i, para in enumerate(doc.paragraphs):
        if para._element.xpath('.//w:br[@w:type="page"]'):
            insert_pos = i + 1
            break

    if insert_pos is None:
        insert_pos = 0

    # 插入目录标题
    toc_title = doc.paragraphs[insert_pos].insert_paragraph_before('目录')
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.style = 'Heading 1'

    # 插入目录域
    toc_para = doc.paragraphs[insert_pos + 1].insert_paragraph_before()
    run = toc_para.add_run()

    # 创建 TOC 域
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._element.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar2)

    run._element.append(OxmlElement('w:t'))

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar3)

    # 目录后分页
    from docx.enum.text import WD_BREAK
    doc.paragraphs[insert_pos + 2].insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)

def adjust_image_size(doc, max_height_cm=24):
    """调整图片大小，限制最大高度为 max_height_cm 厘米"""
    # 1 cm = 360000 EMU (English Metric Units)
    max_height_emu = max_height_cm * 360000

    # 遍历文档中的所有图片
    for shape in doc.inline_shapes:
        # 检查图片高度
        if shape.height > max_height_emu:
            # 计算缩放比例
            scale = max_height_emu / shape.height
            # 按比例调整宽度和高度
            shape.height = int(max_height_emu)
            shape.width = int(shape.width * scale)

def merge_table_cells(table):
    """合并表格中标记为 ^ 的单元格"""
    # 遍历每一列
    for col_idx in range(len(table.columns)):
        merge_start = None
        cells_to_merge = []

        # 遍历每一行
        for row_idx in range(len(table.rows)):
            cell = table.cell(row_idx, col_idx)
            cell_text = cell.text.strip()

            if cell_text == '^':
                # 标记需要合并的单元格
                if merge_start is None:
                    merge_start = row_idx - 1
                cells_to_merge.append(row_idx)
            else:
                # 执行合并
                if merge_start is not None and cells_to_merge:
                    start_cell = table.cell(merge_start, col_idx)
                    for merge_row in cells_to_merge:
                        end_cell = table.cell(merge_row, col_idx)
                        # 清除 ^ 符号
                        for paragraph in end_cell.paragraphs:
                            paragraph.clear()
                        # 合并单元格
                        start_cell.merge(end_cell)

                # 重置
                merge_start = None
                cells_to_merge = []

        # 处理最后一组合并
        if merge_start is not None and cells_to_merge:
            start_cell = table.cell(merge_start, col_idx)
            for merge_row in cells_to_merge:
                end_cell = table.cell(merge_row, col_idx)
                # 清除 ^ 符号
                for paragraph in end_cell.paragraphs:
                    paragraph.clear()
                # 合并单元格
                start_cell.merge(end_cell)

def main(doc_path):
    """添加页码、更新目录、设置表格边框"""
    doc = Document(doc_path)

    # 为所有节添加页码
    for section in doc.sections:
        add_page_number(section)

    # 替换目录标题为中文（保留 TOC 域代码）
    replace_toc_title(doc)

    # 删除手动添加的Markdown目录
    remove_manual_toc(doc)

    # 合并表格单元格并设置边框
    merged_count = 0
    for table in doc.tables:
        # 先合并单元格
        merge_table_cells(table)
        merged_count += 1
        # 再设置边框
        set_table_borders(table)

    # 调整图片大小
    adjust_image_size(doc, max_height_cm=24)

    # 保存文档
    doc.save(doc_path)
    print(f'✅ 页码已添加')
    print(f'✅ 目录标题已替换为中文')
    print(f'✅ 表格单元格已合并（共 {merged_count} 个表格）')
    print(f'✅ 表格边框已应用（共 {len(doc.tables)} 个表格）')
    print(f'✅ 图片大小已调整（最大高度：24cm）')

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python3 adjust-cover-page.py <Word文件>')
        sys.exit(1)

    main(sys.argv[1])
