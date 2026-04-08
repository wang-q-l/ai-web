"""
按照参考文档（销售华东分公司作业票系统升级国标建设方案v2.docx）的格式注入样式：
- 正文：仿宋_GB2312，16pt
- H1：仿宋，15pt，加粗，行距28.8pt，段前5pt，段后4.5pt
- H2：仿宋/Arial，14pt，加粗，行距20.6pt，段前1pt，段后1pt
- H3：仿宋，14pt，加粗，行距20.6pt，段前1pt，段后1pt
- H4：仿宋/Arial，12pt，加粗，行距18.6pt，段前2pt，段后2.5pt
- 页面：A4，左右3.17cm，上下2.54cm
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import zipfile, os, re

DOCX_PATH = "/Users/mac/Desktop/Micon/micon-ai/jssy-anquanfx/.claude/skills/req-doc/assets/reference.docx"

doc = Document(DOCX_PATH)


def set_cn_font(style, ascii_font, ea_font):
    """设置样式的西文和中文字体"""
    rPr = style.font._element
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), ea_font)
    # 清除主题字体引用，防止被主题覆盖
    for attr in [qn('w:asciiTheme'), qn('w:hAnsiTheme'), qn('w:eastAsiaTheme')]:
        if rFonts.get(attr):
            del rFonts.attrib[attr]


def set_line_spacing_auto(style, line_twips):
    """设置 auto 行距（单位 twips，1pt=20twips）"""
    pPr = style.element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        style.element.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(line_twips))
    spacing.set(qn('w:lineRule'), 'auto')


# ── 正文 Normal：仿宋_GB2312，16pt ───────────────────────
normal = doc.styles['Normal']
normal.font.name = '仿宋_GB2312'
normal.font.size = Pt(16)
normal.font.bold = False
set_cn_font(normal, '仿宋_GB2312', '仿宋_GB2312')

# ── Heading 1：仿宋，15pt，加粗，行距28.8pt ──────────────
h1 = doc.styles['Heading 1']
h1.font.name = '仿宋_GB2312'
h1.font.size = Pt(15)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
h1.paragraph_format.space_before = Pt(5)
h1.paragraph_format.space_after = Pt(4.5)
h1.paragraph_format.keep_with_next = True
set_cn_font(h1, '仿宋_GB2312', '仿宋')
set_line_spacing_auto(h1, 576)  # 28.8pt * 20

# ── Heading 2：Arial/仿宋，14pt，加粗，行距20.6pt ─────────
h2 = doc.styles['Heading 2']
h2.font.name = 'Arial'
h2.font.size = Pt(14)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
h2.paragraph_format.space_before = Pt(1)
h2.paragraph_format.space_after = Pt(1)
h2.paragraph_format.keep_with_next = True
set_cn_font(h2, 'Arial', '仿宋')
set_line_spacing_auto(h2, 413)  # 20.65pt * 20

# ── Heading 3：仿宋，14pt，加粗，行距20.6pt ──────────────
h3 = doc.styles['Heading 3']
h3.font.name = '仿宋_GB2312'
h3.font.size = Pt(14)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
h3.paragraph_format.space_before = Pt(1)
h3.paragraph_format.space_after = Pt(1)
h3.paragraph_format.keep_with_next = True
set_cn_font(h3, '仿宋_GB2312', '仿宋')
set_line_spacing_auto(h3, 413)

# ── Heading 4：Arial/仿宋，12pt，加粗，行距18.6pt ─────────
try:
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Arial'
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h4.paragraph_format.space_before = Pt(2)
    h4.paragraph_format.space_after = Pt(2.5)
    h4.paragraph_format.keep_with_next = True
    set_cn_font(h4, 'Arial', '仿宋')
    set_line_spacing_auto(h4, 372)  # 18.6pt * 20
except:
    pass

# ── Table 样式 ───────────────────────────────────────────
try:
    tbl_style = doc.styles['Table Grid']
    tbl_style.font.size = Pt(12)
    set_cn_font(tbl_style, '仿宋_GB2312', '仿宋_GB2312')
except:
    pass

# ── 页面设置：A4，页边距 ──────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3.17)
section.right_margin  = Cm(3.17)
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)

# ── 页眉（空，无下划线）──────────────────────────────────
header = section.header
if not header.paragraphs:
    header.add_paragraph()
hp = header.paragraphs[0]
hp.clear()

# ── 页脚（页码居中）──────────────────────────────────────
footer = section.footer
if not footer.paragraphs:
    footer.add_paragraph()
fp = footer.paragraphs[0]
fp.clear()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

run_pre = fp.add_run('第 ')
run_pre.font.size = Pt(10.5)
run_pre.font.name = '仿宋_GB2312'

fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.text = 'PAGE'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
r_page = OxmlElement('w:r')
r_page.append(fldChar1)
r_page.append(instrText)
r_page.append(fldChar2)
fp._p.append(r_page)

run_mid = fp.add_run(' 页，共 ')
run_mid.font.size = Pt(10.5)
run_mid.font.name = '仿宋_GB2312'

fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'begin')
instrText2 = OxmlElement('w:instrText')
instrText2.text = 'NUMPAGES'
fldChar4 = OxmlElement('w:fldChar')
fldChar4.set(qn('w:fldCharType'), 'end')
r_total = OxmlElement('w:r')
r_total.append(fldChar3)
r_total.append(instrText2)
r_total.append(fldChar4)
fp._p.append(r_total)

run_suf = fp.add_run(' 页')
run_suf.font.size = Pt(10.5)
run_suf.font.name = '仿宋_GB2312'

doc.save(DOCX_PATH)

# ── 修改主题字体（全部换成仿宋）────────────────────────────
tmp_path = DOCX_PATH + ".tmp"
with zipfile.ZipFile(DOCX_PATH, 'r') as zin:
    with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/theme/theme1.xml':
                xml = data.decode('utf-8')
                xml = re.sub(r'(<a:majorFont[^>]*>.*?<a:ea\s+typeface=")[^"]*(")', r'\g<1>仿宋_GB2312\g<2>', xml, flags=re.DOTALL)
                xml = re.sub(r'(<a:minorFont[^>]*>.*?<a:ea\s+typeface=")[^"]*(")', r'\g<1>仿宋_GB2312\g<2>', xml, flags=re.DOTALL)
                xml = re.sub(r'(<a:majorFont[^>]*>.*?<a:font\s+script="Hans"\s+typeface=")[^"]*(")', r'\g<1>仿宋_GB2312\g<2>', xml, flags=re.DOTALL)
                xml = re.sub(r'(<a:minorFont[^>]*>.*?<a:font\s+script="Hans"\s+typeface=")[^"]*(")', r'\g<1>仿宋_GB2312\g<2>', xml, flags=re.DOTALL)
                data = xml.encode('utf-8')
            zout.writestr(item, data)

os.replace(tmp_path, DOCX_PATH)
print("✅ reference.docx 样式注入完成（仿宋_GB2312 格式）")
