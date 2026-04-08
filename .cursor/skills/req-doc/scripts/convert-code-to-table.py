#!/usr/bin/env python3
"""
将代码块段落转换为单元格表格
"""
import sys
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def convert_code_blocks_to_tables(doc_path):
    """将所有代码块段落转换为单元格表格"""
    doc = Document(doc_path)

    # 收集所有代码块段落
    code_blocks = []
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == 'Source Code':
            code_blocks.append((i, para))

    print(f'找到 {len(code_blocks)} 个代码块')

    # 从后往前处理，避免索引变化
    for idx, para in reversed(code_blocks):
        # 获取段落内容
        text = para.text
        if not text.strip():
            continue

        # 在段落位置插入表格
        para_element = para._element
        parent = para_element.getparent()
        para_index = parent.index(para_element)

        # 创建表格（1行1列）
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]

        # 设置单元格内容
        cell.text = text

        # 设置单元格字体和对齐方式
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        for paragraph in cell.paragraphs:
            # 设置左对齐
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # 取消首行缩进
            paragraph.paragraph_format.first_line_indent = Pt(0)

            for run in paragraph.runs:
                run.font.name = '仿宋_GB2312'
                run.font.size = Pt(10)

                # 设置中文字体
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

        # 设置单元格背景色
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F5F5F5')
        tcPr.append(shd)

        # 设置表格边框
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), 'D4D4D4')
            tblBorders.append(border)
        tblPr.append(tblBorders)

        # 设置单元格内边距
        tblCellMar = OxmlElement('w:tblCellMar')
        for margin_name in ['top', 'bottom', 'left', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '100')
            margin.set(qn('w:type'), 'dxa')
            tblCellMar.append(margin)
        tblPr.append(tblCellMar)

        # 将表格插入到段落位置
        parent.insert(para_index, table._element)

        # 删除原段落
        parent.remove(para_element)

    doc.save(doc_path)
    print(f'✅ 代码块已转换为表格（共 {len(code_blocks)} 个）')

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python3 convert-code-to-table.py <文档路径>')
        sys.exit(1)

    convert_code_blocks_to_tables(sys.argv[1])
